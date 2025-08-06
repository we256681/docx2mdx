#!/usr/bin/env python3

import os
from docx import Document
import markdown
import sys
import re
import yaml


def read_document(docx_path):
    """Read a Word document with error handling."""
    # Check if the file exists
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: File '{docx_path}' not found.")

    # Check if the file is a .docx file
    if not docx_path.lower().endswith('.docx'):
        raise ValueError(f"Error: File '{docx_path}' is not a .docx file.")

    # Check if the file is not empty
    if os.path.getsize(docx_path) == 0:
        raise ValueError(f"Error: File '{docx_path}' is empty.")

    try:
        print(f'Opening {docx_path}')
        doc = Document(docx_path)
        print(f"Successfully opened '{docx_path}'.")
        return doc
    except Exception as e:
        raise IOError(f"Error reading file '{docx_path}': {e}")



def parse_media_url(vals):
    '''Ensure that the media link is correctly formatted 
    and outputs are always saved to media directory.
    '''
    if 'media' in vals:
        # Regex to find if the path already starts with './media/' followed by non-space characters.
        # This is to ensure we don't prepend './media/' if it's already correctly formatted.
        pattern = r'(\./media/[^\s]+)'
        match = re.search(pattern, vals)
        if match:
            vals = match.group(1)  # Use the already correctly formatted path.
        else:
            # If not, construct the path assuming the value after ': ' is the filename.
            vals = f"./media/{vals.split(': ')[-1]}"
    return vals

def parse_media_alt_text(all_text):
    """
    Parses a block of text to extract media-related information like alt text,
    author name, and author URL using regular expressions.
    """
    output = []
    # Regex to capture the value after "Image text (alt): " up to the newline.
    match_alt_text = re.search(r'Image text \(alt\): (.*?)\n', all_text, re.IGNORECASE)
    # Regex to capture the value after "Author name: " up to the newline.
    match_author_name = re.search(r'Author name: (.*?)\n',all_text, re.IGNORECASE)
    # Regex to capture the value after "Author URL: " (possibly with spaces) up to the end of the line.
    match_author_URL = re.search(r'Author URL:\s*(.*)', all_text, re.IGNORECASE)

    out_names = ['main_fig_alt_text', 'main_fig_author_name', 'main_fig_author_URL']
    out_data = [match_alt_text, match_author_name, match_author_URL]

    for idx,match_ in enumerate(zip(out_names, out_data)):
        if match_[1]: # If regex found a match
            output.append({match_[0]: match_[1].group(1)})
        else: # If no match, provide a default value
            output.append({match_[0]: "Data not provided"})
    return output


def parse_tag_information(all_text):
    """
    Parses a block of text to extract tag-related information like topic, subtopic,
    source, and product type using regular expressions.
    """
    output = []
    # Regex to capture the value after "Topic: " up to the newline.
    match_topic_values = re.search(r'Topic: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to capture the value after "Subtopic: " up to the newline.
    match_subtopic_values = re.search(r'Subtopic: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to capture the value after "Source: " up to the newline.
    match_source = re.search(r'Source: (.*?)\n',all_text, re.IGNORECASE)
    # Regex to capture the value after "Product Type: " (possibly with spaces) up to the end of the line.
    match_prod_type = re.search(r'Product Type:\s*(.*)', all_text, re.IGNORECASE)

    out_names = ['topic', 'subtopic', 'source', 'product_type']
    out_data = [match_topic_values, match_subtopic_values, match_source, match_prod_type]

    for idx,match_ in enumerate(zip(out_names, out_data)):
        if match_[1]: # If regex found a match
            output.append({match_[0]: match_[1].group(1)})
        else: # If no match, provide a default value
            output.append({match_[0]: "Data not provided"})
    return output


def parse_layer_information(all_text):
    """
    Parses a block of text to extract detailed information for multiple layers.
    Each piece of information for a layer is expected to be on a new line,
    prefixed by a label (e.g., "Layer name: ...").
    """
    # Regex to find "Layer name: " followed by any characters until a newline.
    layer_name = re.findall(r'Layer name: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "stacCol: " followed by any characters until a newline.
    stacCol = re.findall(r'stacCol: (.*?)\n',all_text, re.IGNORECASE)
    # Regex to find "Layer id: " followed by any characters until a newline.
    layer_id = re.findall(r'Layer id: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Layer description: " followed by any characters until a newline.
    layer_description = re.findall(r'Layer description: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Units: " followed by any characters until a newline.
    unit = re.findall(r'Units: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Color ramp description: " followed by any characters until a newline.
    color_ramp_description = re.findall(r'Color ramp description: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find content within square brackets `[...]`, used for color stops.
    color_stops = re.findall(r'\[([^\]]+)\]', all_text, re.IGNORECASE)
    # Regex to find "Data format: " followed by any characters until a newline.
    data_format = re.findall(r'Data format: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Projection: " followed by any characters until a newline.
    projection = re.findall(r'Projection: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Legend type: " followed by any characters until a newline.
    legend_type = re.findall(r'Legend type: (.*?)\n', all_text, re.IGNORECASE) # Note: Duplicated in original, kept for now
    # Regex to find "Legend minimum: " followed by any characters until a newline.
    legend_min = re.findall(r'Legend minimum: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Legend maximum: " followed by any characters until a newline.
    legend_max = re.findall(r'Legend maximum: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Legend type: " followed by any characters until a newline. (This is a duplicate regex from above)
    legend_type = re.findall(r'Legend type: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Colormap name: " followed by any characters until a newline.
    colormap_name = re.findall(r'Colormap name: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Resampling: " followed by any characters until a newline.
    resampling = re.findall(r'Resampling: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Rescale minimum: " followed by any characters until a newline.
    rescale_min = re.findall(r'Rescale minimum: (.*?)\n', all_text, re.IGNORECASE)
    # Regex to find "Rescale maximum: " followed by any characters until a newline.
    rescale_max = re.findall(r'Rescale maximum: (.*?)\n', all_text, re.IGNORECASE)

    # Clean up color_stops: each match is a string of comma-separated items.
    # This loop processes each string, splits it, strips individual items,
    # and removes any empty items resulting from stray commas or spaces.
    final_color_groups = []
    for match in color_stops:
        items = [item.strip(" '\n") for item in match.split(",")]
        # Remove empty items that might result from trailing commas or multiple commas.
        items = [item for item in items if item]
        final_color_groups.append(items)

    num_layers = len(layer_name) # Assuming layer_name is a reliable indicator of the number of layers.

    # Define the keys for layer data and the corresponding lists of parsed values.
    # This structure helps in iterating and assigning values to each layer's dictionary.
    out_names = ['layer_name', 'stacCol', 'layer_id', 'layer_description', 'units', 
                 'color_ramp_description', 'color_stops', 'data_format','projection',
                 'legend_minimum','legend_maximum','legend_type','colormap_name',
                 'resampling','rescale_min','rescale_max']
    out_data = [layer_name, stacCol, layer_id, layer_description, unit,
                 color_ramp_description, final_color_groups, data_format, projection, 
                 legend_min, legend_max, legend_type, colormap_name,
                 resampling, rescale_min, rescale_max]

    output = []
    # Iterate based on the number of layer names found.
    for i in range(num_layers):
        layer_data = {}  # Create a dictionary for the current layer.
        # Populate layer_data using out_names and out_data.
        for idx, match_info in enumerate(zip(out_names, out_data)):
            try:
                # Assign the i-th element of the current data list (e.g., layer_name[i], stacCol[i])
                # to the corresponding key (e.g., layer_name0, stacCol0).
                layer_data[f'{match_info[0]}{i}'] = match_info[1][i]
            except IndexError:
                # If a particular piece of data is missing for this layer (e.g., no rescale_max for layer i),
                # an IndexError might occur. Pass silently to allow partial data.
                pass
        output.append(layer_data)

    return output


def table_0_info(row, header, extracted_data):
    """
    Extracts information from the first table (table_0) of the document.
    This table usually contains general metadata, media links, tags, and layer information.
    """
    # Special handling for 'media' field.
    if header == "media":
        all_text = parse_media_alt_text(row.cells[1].text.strip()) # Extract alt text, author, etc.
        vals = parse_media_url(row.cells[1].text.strip().split("\n")[0]) # Extract and format the main media URL.
        
        if header not in extracted_data:
            extracted_data[header] = []

        extracted_data[header].append({'main_media_image': check_value_string_length(vals)})
        for idx, val in enumerate(all_text): # Append other media details like alt text.
            extracted_data[header].append({next(iter(val.keys())):val[next(iter(val.keys()))]})
    # Special handling for 'tags' field.
    elif header == 'tags':
        all_text = parse_tag_information(row.cells[1].text.strip()) # Extract topics, source, etc.
        if header not in extracted_data:
            extracted_data[header] = []

        for idx, val in enumerate(all_text): # Append tag details.
            extracted_data[header].append({next(iter(val.keys())):val[next(iter(val.keys()))]})
    # Special handling for 'layers' field.
    elif header == 'layers':
        print(f'Parsing layer information.')
        all_text = parse_layer_information(row.cells[1].text.strip()) # Extract detailed layer data.

        for i in range(len(all_text)): # Append each layer's data.
            if header not in extracted_data:
                extracted_data[f'{header}'] = []
            extracted_data[f'{header}'].append({f'Layer{i}':all_text[i]})
    # Default handling for other fields.
    else:
        vals = row.cells[1].text.strip().split("\n")[0]
        extracted_data[header] = check_value_string_length(vals)

    return extracted_data

def check_value_string_length(vals):
    """If string is empty, then replace the value with 'Data not provided'."""
    return vals if len(vals) >=1 else "Data not provided"

def parse_table_value_content(row,header,table_0,table_1):
    """
    Parses content from table_1, which typically contains specific dataset attributes
    like content source, temporal extent, etc., often formatted as "Value: ...".
    """
    all_text = check_value_string_length(row.cells[1].text.strip())
    
    # Regex to extract content after "Value: " (case-insensitive).
    value_regex = r'(?<=Value:\s)(.*)'

    if header == 'content_source':
        # For content_source, if the value is 'null', try to get it from table_0's tags.
        # This provides a fallback mechanism.
        extracted_values = re.findall(value_regex, all_text, re.IGNORECASE)
        if extracted_values and extracted_values[0].lower() == 'null':
            source_values = [item['source'] for item in table_0.get('tags', []) if 'source' in item]
            table_1[header] = source_values if source_values else ["Data not provided"]
        else:
            table_1[header] = extracted_values if extracted_values else ["Data not provided"]
    elif header == 'temporal_extent':
        # Regex to find "Start: MM/DD/YYYY" or "End: MM/DD/YYYY" and capture the dates.
        # It allows for variations in spacing.
        match_value_names = re.findall(r'Start:\s*(\d{2}/\d{2}/\d{4})|End:\s*(\d{2}/\d{2}/\d{4})', all_text, re.IGNORECASE)
        # Extract first matched start date and end date.
        start_value = next((match[0] for match in match_value_names if match[0]), None)
        end_value = next((match[1] for match in match_value_names if match[1]), None)
        table_1['start_temporal_extent'] = start_value if start_value else "Data not provided"
        table_1['end_temporal_extent'] = end_value if end_value else "Data not provided"
    elif header == 'legend_value_range':
         pass # This section was intentionally commented out in the original code.
        # #This will add if all of the layers are identical, but that isn't always the case
        # match_value_names = re.findall(r"Min:\s*([\d.]+)\s*Max:\s*([\d.]+)\s*Type:\s*(\w+)", all_text)
        # min_value = next((match[0] for match in match_value_names if match[0]), None)
        # max_value = next((match[1] for match in match_value_names if match[1]), None)
        # legend_type = next((match[2] for match in match_value_names if match[2]), None)
        # table_1['legend_min'] = min_value
        # table_1['legend_max'] = max_value
        # table_1['legend_type'] = legend_type
    else:
        # Default extraction for other headers in this table.
        match_value_names = re.findall(value_regex, all_text, re.IGNORECASE)
        table_1[header] = match_value_names if match_value_names else ["Data not provided"]
    return table_1

def parse_additional_table_info(row, header, table_2):
    """
    Parses content from table_2, which may contain additional structured information
    formatted as "Header: ... Value: ...".
    """
    all_text = check_value_string_length(row.cells[1].text.strip())
    # Regex to capture content after "Header: " and "Value: ", allowing for newlines between them.
    # re.DOTALL makes '.' match newlines as well.
    match = re.search(r'Header:\s*(.*?)\s*\n+\s*Value:\s*(.*)', all_text, re.DOTALL)
    
    if match:
        header_ = match.group(1).strip() # Captured header content.
        value = match.group(2).strip()   # Captured value content.
        if len(header_) != 0 and len(value) != 0: # Ensure both are non-empty.
            table_2[header] = [] # Use the original row's first cell as the main key.
            table_2[header].append({header_: value}) # Store as a list containing a dict.
    return table_2


def extract_table_info_from_docx(doc):
    """Extract information from the tables in the Word document."""
    table_0 = {} # Stores general metadata, media, tags, layers.
    table_1 = {} # Stores specific dataset attributes (e.g., temporal extent, source).
    table_2 = {} # Stores additional structured information.
    
    # Process lists (bulleted or numbered)
    # Iterate through all tables in the document.
    for iTable,table in enumerate(doc.tables):
        # Iterate through all rows in the current table.
        for iRow,row in enumerate(table.rows):
            # Get header from the first cell of the row, convert to lowercase, and take the first line.
            header = row.cells[0].text.strip().lower().split("\n")[0]
            
            # Skip rows where both header and the second cell's content are empty.
            if len(header) == 0 and len(row.cells[1].text.strip()) == 0:
                continue
            # If header is empty but there's content, it's an invalid format.
            elif len(header) == 0:
                raise ValueError("Header is empty. Please check the template document and ensure all expected headers are present and not empty.")
            # Process valid rows based on table index.
            else:
                if iTable == 0: # First table (general info, media, tags, layers)
                    table_0 = table_0_info(row, header, table_0)
                elif iTable == 1: # Second table (dataset attributes)
                    table_1 = parse_table_value_content(row,header,table_0,table_1)
                elif iTable ==2: # Third table (additional info)
                    table_2 = parse_additional_table_info(row, header, table_2)
    return table_0, table_1, table_2


def extract_headers_and_paragraphs(doc):
    """
    Extracts content from paragraphs in the document, attempting to group
    paragraphs under the nearest preceding header.
    Headers are identified by style (e.g., "Heading 1") or bold formatting.
    """
    output = {}
    current_header = None
    current_content = []

    for para in doc.paragraphs:
        # Check if paragraph style name starts with 'Heading' or if any run in it is bold.
        # This is a common way to identify headers in DOCX files.
        is_header = para.style.name.startswith('Heading') or any(run.bold for run in para.runs)

        if is_header:
            if current_header:
                # If there's a pending header, save its accumulated content to the output dictionary.
                output[current_header] = "\n".join(current_content).strip()
                current_content = []  # Reset content list for the new header.
            current_header = para.text.strip() # Set the new current header.
        elif current_header:
            # If it's a content paragraph and a header is active, append its text.
            current_content.append(para.text.strip())

    # After the loop, save any remaining content for the last header.
    if current_header and current_content:
        output[current_header] = "\n".join(current_content).strip()

    return output


def retrieve_all_docx_data(docx_path):
    try:
        doc = read_document(docx_path)
    except Exception as e:
        print(e)
    # Extract the text from the DOCX
    table_0, table_1, table_2 = extract_table_info_from_docx(doc)
    content = extract_headers_and_paragraphs(doc)

    return table_0, table_1, table_2, content



